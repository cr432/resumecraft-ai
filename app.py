import os
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
import re

CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY")
CLAUDE_API_URL = "https://api.anthropic.com/v1/messages"
CLAUDE_MODEL = "claude-3-opus-20240229"  # or your preferred model

app = FastAPI(title="Claude Resume Tailor")

# Add CORS middleware to allow frontend to connect
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def create_improved_docx(text, output_path):
    doc = Document()
    
    # Add styles to document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Parse and create sections
    current_section = None
    in_bullet_list = False
    bullet_items = []
    
    lines = text.split('\n')
    
    for line in lines:
        # Skip empty lines
        if not line.strip():
            continue
        
        # Check for section headers (all caps or [SECTION] format)
        if re.match(r'^\[SECTION\]|\[HEADING\]', line, re.IGNORECASE):
            # Add any pending bullet list items
            if in_bullet_list and bullet_items:
                para = doc.add_paragraph()
                for item in bullet_items:
                    bullet_para = doc.add_paragraph(style='List Bullet')
                    bullet_para.add_run(item)
                bullet_items = []
                in_bullet_list = False
            
            # Add section header
            clean_heading = re.sub(r'^\[SECTION\]|\[HEADING\]', '', line, flags=re.IGNORECASE).strip()
            heading = doc.add_heading(clean_heading, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            current_section = clean_heading
            
        # Check for subsection headers
        elif re.match(r'^\[SUBSECTION\]|\[SUBHEADING\]', line, re.IGNORECASE):
            # Add any pending bullet list items
            if in_bullet_list and bullet_items:
                para = doc.add_paragraph()
                for item in bullet_items:
                    bullet_para = doc.add_paragraph(style='List Bullet')
                    bullet_para.add_run(item)
                bullet_items = []
                in_bullet_list = False
            
            clean_subheading = re.sub(r'^\[SUBSECTION\]|\[SUBHEADING\]', '', line, flags=re.IGNORECASE).strip()
            subheading = doc.add_heading(clean_subheading, level=2)
            subheading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
        # Check for bullet points
        elif line.strip().startswith('•') or line.strip().startswith('-') or line.strip().startswith('*') or re.match(r'^\[BULLET\]', line, re.IGNORECASE):
            clean_bullet = re.sub(r'^[•\-\*]|\[BULLET\]', '', line, flags=re.IGNORECASE).strip()
            bullet_items.append(clean_bullet)
            in_bullet_list = True
            
        # Check for bold text
        elif re.match(r'^\[BOLD\]', line, re.IGNORECASE):
            clean_text = re.sub(r'^\[BOLD\]', '', line, flags=re.IGNORECASE).strip()
            para = doc.add_paragraph()
            para.add_run(clean_text).bold = True
            
        # Regular text
        else:
            # Add any pending bullet list items
            if in_bullet_list and bullet_items:
                for item in bullet_items:
                    bullet_para = doc.add_paragraph(style='List Bullet')
                    bullet_para.add_run(item)
                bullet_items = []
                in_bullet_list = False
            
            # Check if it might be a section header (all caps)
            if line.isupper() and len(line) < 50:
                heading = doc.add_heading(line, level=1)
                heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                current_section = line
            else:
                para = doc.add_paragraph(line)
    
    # Add any remaining bullet items
    if in_bullet_list and bullet_items:
        for item in bullet_items:
            bullet_para = doc.add_paragraph(style='List Bullet')
            bullet_para.add_run(item)
    
    doc.save(output_path)

@app.post("/tailor-resume")
async def tailor_resume(
    resume: UploadFile = File(...),
    job_description: str = Form(...)
):
    if not resume.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")
    if not CLAUDE_API_KEY:
        raise HTTPException(status_code=500, detail="Claude API key not set in environment.")

    # Save uploaded resume to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_resume:
        temp_resume.write(await resume.read())
        temp_resume_path = temp_resume.name

    # Extract text from resume
    resume_text = extract_text_from_docx(temp_resume_path)

    # Prepare prompt for Claude with structured output instructions
    prompt = """
You are an expert resume writer with years of experience helping job seekers create impressive resumes. Given the following resume and job description, transform the content while PRESERVING THE ORIGINAL RESUME'S FORMATTING AND STYLE.

Important instructions:
1. Keep the same overall structure, section organization, and formatting style as the original resume
2. Follow the same pattern of bullet points, headers, and text formatting
3. Preserve the original resume's tone and presentation style
4. Generate better bullet points that are more relevant to the job description
5. Highlight skills and experiences that match the job requirements
6. Adjust content to better position the candidate for this specific role
7. Remove or reduce emphasis on irrelevant experiences
8. Study the layout and style of the original resume carefully, and maintain that same overall look and feel. The goal is to create a tailored resume that doesn't look completely different from the original.

Return ONLY the improved resume text with formatting indicators as described above. Do not include any other text.

Job Description:
"""
    prompt += job_description
    prompt += "\n\nResume:\n"
    prompt += resume_text

    # Call Claude API
    headers = {
        "x-api-key": CLAUDE_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json"
    }
    data = {
        "model": CLAUDE_MODEL,
        "max_tokens": 4096,
        "messages": [
            {"role": "user", "content": prompt}
        ]
    }
    response = requests.post(CLAUDE_API_URL, headers=headers, json=data)
    if response.status_code != 200:
        os.remove(temp_resume_path)
        raise HTTPException(status_code=500, detail=f"Claude API error: {response.text}")
    improved_text = response.json()["content"][0]["text"]

    # Create improved docx
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_out:
        create_improved_docx(improved_text, temp_out.name)
        output_path = temp_out.name

    os.remove(temp_resume_path)
    return FileResponse(output_path, filename=f"tailored_{resume.filename}")

@app.get("/")
def root():
    with open("client.html", "r") as f:
        html_content = f.read()
    return HTMLResponse(content=html_content) 